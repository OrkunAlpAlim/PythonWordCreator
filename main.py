import subprocess
import sys
import os

def check_install(package):
   # Kütüphane yüklü mü kontrol eder, değilse yükler
    try:
        __import__(package)
        print(f"{package} is already installed.")
    except ImportError:
        print(f"{package} is not installed. Installing...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

def main():
    # Önce python-docx kütüphanesinin yüklü olup olmadığını kontrol et
    check_install("python-docx")

    # Gerekli modül yüklendikten sonra import edilir
    from docx import Document
    from docx.shared import RGBColor, Pt
    
    # Dosya adı ve yolu
    doc_name = "example.docx"
    doc_path = ""  # Bu kısmı kendi yolunuza göre ayarlayın
    
    # Dosya yolunun olup olmadığını kontrol et
    if not os.path.exists(doc_path):
       doc_path = "C:\\Users\\OrkunAlpALİM\\Desktop"  # Eğer path yoksa default path atar

    # Yeni bir Word belgesi oluştur
    doc = Document()

    # Paragraf ekle
    paragraph = doc.add_paragraph()

    # İlk metin (run) - Mavi renkte ve kalın
    part1 = paragraph.add_run("The first run of the paragraph. \n")
    part1.font.color.rgb = RGBColor(0, 0, 255)
    part1.font.bold = True

    # İkinci metin (run) - Kırmızı renkte, italik ve küçük boyutta
    run2 = paragraph.add_run("The second run of the paragraph.")
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(255, 0, 0)

    # Belgeyi belirlenen yola kaydet
    full_doc_path = os.path.join(doc_path, doc_name)
    doc.save(full_doc_path)  # Dosya yoluna kaydet

    # Dosyanın başarıyla oluşturulup oluşturulmadığını kontrol et
    if os.path.exists(full_doc_path):
        print(f"The Word file was created successfully, and saved as {full_doc_path}")
    else:
        print("An issue occurred while creating the Word file")
    
    # Belge yolunu return et
    return full_doc_path

# Ana fonksiyonu çalıştır
if __name__ == "__main__":
    doc_path = main()
    if doc_path:
        print(f"Document path returned: {doc_path}")
    else:
        print("No document created.")
