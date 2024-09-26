import subprocess
import sys
import os
from docx import Document
from docx.shared import RGBColor, Pt

def check_install(package):
    # Kütüphane yüklü değilse yükle
    try:
        __import__(package)
        print(f"{package} is already installed.")
    except ImportError:
        print(f"{package} is not installed. Installing...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

def main():
    # Gerekli kütüphanenin parametresi gönder
    check_install("python-docx")
    
    # Yeni bir Word belgesi oluştur
    doc = Document()

    # Dosya adı ve yolu
    doc_name = "example.docx"
    doc_path = ""  # Bu kısmı kendi yolunuza göre ayarlayın
    
    # Paragraf ekle
    paragraph = doc.add_paragraph()

    # İlk metin (run) - Mavi renkte ve kalın
    part1 = paragraph.add_run("The first run of the paragraph. \n")
    part1.font.color.rgb = RGBColor(0, 0, 255)
    part1.font.bold = True

    # İkinci metin (run) - Kırmızı renkte, italik ve küçük boyutta
    run2 = paragraph.add_run("The second run of the paragraph.")
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(255, 0, 0)

    # Belgeyi belirlenen yola kaydet (path null ise default path belirlenir)
    if not doc_path:
        doc_path = "/Users/orkunalpalim/Desktop/"
    full_doc_path = os.path.join(doc_path, doc_name)
    doc.save(full_doc_path)  # Dosya yoluna kaydet

    # Dosyanın başarıyla oluşturulup oluşturulmadığını kontrol et
    if os.path.exists(full_doc_path):
        print(f"The Word file was created successfully, and saved as {full_doc_path}")
    else:
        print("An issue occurred while creating the Word file")
    
    # Belge yolunu return et
    return full_doc_path

# Ana fonksiyonu çalıştır
if __name__ == "__main__":
    doc_path = main()
