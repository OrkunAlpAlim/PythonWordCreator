from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE, WD_COLOR_INDEX
from docx.oxml import OxmlElement

# Word dosyası oluştur
doc = Document()

# Başlık ekle
doc.add_heading('python-docx Kütüphanesi - Tüm Font Özelliklerinin Kullanımı', level=0)

# Font Boyutu
doc.add_heading('1. Font Boyutu', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin 20 punto boyutundadır ve ')
run.font.size = Pt(20)
run2 = paragraph.add_run('20 punto boyutu kullanılarak yazılmıştır.')
run2.font.size = Pt(20)

# Yazı Tipi
doc.add_heading('2. Yazı Tipi', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin Arial yazı tipi ile ')
run.font.name = 'Arial'
run2 = paragraph.add_run('yazılmıştır.')
run2.font.name = 'Arial'

# Kalın Yazı
doc.add_heading('3. Kalın Yazı', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin kalın yazı ile ')
run.font.bold = True
run2 = paragraph.add_run('yazılmıştır.')
run2.font.bold = True

# İtalik Yazı
doc.add_heading('4. İtalik Yazı', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin italik yazı ile ')
run.font.italic = True
run2 = paragraph.add_run('yazılmıştır.')
run2.font.italic = True

# Altı Çizili Yazı
doc.add_heading('5. Altı Çizili Yazı', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin altı çizili yazı ile ')
run.font.underline = WD_UNDERLINE.SINGLE
run2 = paragraph.add_run('yazılmıştır.')
run2.font.underline = WD_UNDERLINE.SINGLE

# Yazı Rengi
doc.add_heading('6. Yazı Rengi', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin kırmızı renkte ')
run.font.color.rgb = RGBColor(255, 0, 0)
run2 = paragraph.add_run('yazılmıştır.')
run2.font.color.rgb = RGBColor(255, 0, 0)

# Vurgu Rengi (Highlight)
doc.add_heading('7. Vurgu Rengi', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin sarı renkle vurgulanmıştır.')
run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# Satır Kesmesi (Soft Return)
doc.add_heading('8. Satır Kesmesi', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin satır kesmesi ile')
run._r.append(OxmlElement('w:br'))
run2 = paragraph.add_run(' alt satıra geçmiştir.')

# Büyük Harf Stili (All Caps)
doc.add_heading('9. Büyük Harf Stili', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin BÜYÜK HARF STİLİ ile yazılmıştır.')
run.font.all_caps = True

# Küçük Büyük Harf Stili (Small Caps)
doc.add_heading('10. Küçük Büyük Harf Stili', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin Küçük Büyük Harf Stili ile yazılmıştır.')
run.font.small_caps = True

# Çizik Yazı (Strikethrough)
doc.add_heading('11. Çizik Yazı', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin üstü çizili olarak yazılmıştır.')
run.font.strike = True

# Gizli Metin (Hidden Text)
doc.add_heading('12. Gizli Metin', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('Bu metin gizlenmiştir ve Word üzerinde görünmeyecektir.')
run.font.hidden = True

# Alt Simge ve Üst Simge
doc.add_heading('13. Alt Simge ve Üst Simge', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run('H2O')
run.font.subscript = True  # Alt simge
run2 = paragraph.add_run(' ve ')
run3 = paragraph.add_run('E=mc2')
run3.font.superscript = True  # Üst simge

# Dosyayı kaydet
doc.save('example.docx')
